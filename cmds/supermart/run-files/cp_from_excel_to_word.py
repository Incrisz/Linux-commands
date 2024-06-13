import os
import pandas as pd
from docx import Document

def get_excel_files(directory):
    """Get all Excel files in the specified directory and its subdirectories."""
    excel_files = []
    for root, _, files in os.walk(directory):
        for file in files:
            if file.endswith(('.xlsx', '.xls')):
                excel_files.append(os.path.join(root, file))
    return excel_files

def extract_product_names(excel_file):
    """Extract 'Product Name' column from an Excel file."""
    try:
        df = pd.read_excel(excel_file)
        # Debugging: Print out the column names
        print(f"Columns in {excel_file}: {df.columns.tolist()}")
        
        if 'Product Name' in df.columns:
            return df['Product Name'].dropna().tolist()
        else:
            print(f"'Product Name' column not found in {excel_file}")
    except Exception as e:
        print(f"Error reading {excel_file}: {e}")
    return []

def write_to_word(product_names, output_file):
    """Write product names to a Word document."""
    doc = Document()
    doc.add_heading('Product Names', level=1)
    for name in product_names:
        doc.add_paragraph(name)
    doc.save(output_file)

def main():
    # Specify the directory to scan
    directory = 'C:/Users/USER/Downloads/easelow (1)/supermart'  # Change this to your directory path
    
    # Get all Excel files in the directory
    excel_files = get_excel_files(directory)
    
    # Extract product names from each Excel file
    all_product_names = []
    for excel_file in excel_files:
        product_names = extract_product_names(excel_file)
        all_product_names.extend(product_names)
    
    # Write the product names to a Word document
    output_file = 'market.docx'
    write_to_word(all_product_names, output_file)
    print(f"Product names have been written to {output_file}")

if __name__ == '__main__':
    main()